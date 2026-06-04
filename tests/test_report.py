"""Unit tests for the HTML/PDF report renderer."""
from __future__ import annotations

from pathlib import Path

import pytest

from sespy import data_structure as ds
from sespy.report import render_html, render_pdf

SAMPLE = Path(__file__).resolve().parents[1] / "data" / "sample_ses.json"


@pytest.fixture(scope="module")
def project() -> ds.Project:
    return ds.Project.from_isa(ds.load_sample(SAMPLE), name="Sample SES")


def test_render_html_returns_full_document(project):
    html = render_html(project)
    assert html.startswith("<!DOCTYPE html>")
    assert html.rstrip().endswith("</html>")
    # Project name + a known sample element show up in the report
    assert "Sample SES" in html
    assert "Tourism demand" in html


def test_render_html_includes_summary_counts(project):
    html = render_html(project)
    assert ">17<" in html  # element count
    assert ">20<" in html  # connection count


def test_render_html_includes_loop_classifications(project):
    html = render_html(project)
    # Pills carry the type as their text content; HTML-encoded
    assert "Reinforcing" in html or "reinforcing" in html


def test_render_html_includes_centrality_metrics(project):
    html = render_html(project)
    # Each of the 7 metrics gets its own h3 section
    for metric in ("Degree", "Betweenness", "Closeness", "Eigenvector", "Pagerank"):
        assert metric in html


def test_render_html_includes_leverage_section(project):
    html = render_html(project)
    assert "Leverage points" in html
    assert "Composite score" in html


def test_render_pdf_returns_pdf_bytes(project):
    """PDF round-trip — slow, but cheap insurance against template
    breaking weasyprint's CSS parser. PDF export is an optional feature
    (the `pdf` extra / WeasyPrint + native libs), so skip when it's
    absent rather than fail a core-only install."""
    pytest.importorskip("weasyprint")
    pdf = render_pdf(project)
    assert isinstance(pdf, bytes)
    assert len(pdf) > 5000  # a non-empty PDF is at least a few KB
    assert pdf.startswith(b"%PDF-")


def test_render_docx_returns_valid_docx(project):
    """Word .docx files are zip archives with PK header. Round-trip
    through python-docx to confirm the saved bytes are openable."""
    from io import BytesIO

    from docx import Document  # type: ignore

    from sespy.report import render_docx

    payload = render_docx(project)
    assert isinstance(payload, bytes)
    assert len(payload) > 5000
    # PK is the zip-file magic — .docx files are zip-packaged XML
    assert payload[:2] == b"PK"

    # Round-trip: reopen the bytes and check we can read text from it
    doc = Document(BytesIO(payload))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "SES Toolbox Report" in text
    # Sample element label should appear somewhere (in a table cell)
    full_text = text + "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Tourism demand" in full_text
