"""Report generation — HTML + PDF export of the current project's
ISA data + analysis results.

R counterpart: `modules/prepare_report_module.R` + `functions/report_generation.R`
which use rmarkdown to render `.Rmd` templates. We use Jinja2 + WeasyPrint
which fills the same niche: a parametrised template gets rendered against
project data, then the HTML is either downloaded directly or converted
to PDF.

Why not Quarto? It works for Python, but spinning up a Quarto subprocess
for every download is heavier than rendering a template in-process.
Jinja2 + WeasyPrint runs synchronously in the worker thread.
"""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from jinja2 import Environment, BaseLoader, select_autoescape

from . import network as net_analysis
from .data_structure import IsaData, Project


_TEMPLATE = """\
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <title>{{ project.metadata.name }} — SES Report</title>
  <style>
    @page { size: A4; margin: 18mm 16mm 18mm 16mm; }
    :root {
      --ocean-deep: #0f2744; --ocean-mid: #1a3a5c; --ocean-shallow: #2d5a7b;
      --ocean-surface: #4a90b8; --bio-cyan: #00d4aa; --kelp-green: #52b788;
      --foam: #f8fbfd; --mist: #e8f1f8; --slate: #3d5a73;
      --text: #0a1628; --muted: #506a80;
    }
    body {
      font-family: 'Source Serif 4', Georgia, serif;
      font-size: 11pt; line-height: 1.5; color: var(--text);
      margin: 0; background: white;
    }
    h1, h2, h3 { font-family: 'DM Sans', system-ui, sans-serif;
                 color: var(--ocean-deep); font-weight: 600;
                 letter-spacing: -0.01em; }
    h1 { font-size: 24pt; margin-bottom: 4pt; }
    h2 { font-size: 16pt; margin-top: 24pt;
         border-bottom: 1px solid var(--mist); padding-bottom: 4pt; }
    h3 { font-size: 13pt; margin-top: 16pt; color: var(--ocean-shallow); }
    .meta { color: var(--muted); font-size: 10pt; margin-bottom: 16pt; }
    .summary {
      display: flex; gap: 12pt; margin: 12pt 0 18pt 0;
    }
    .stat {
      background: var(--mist); padding: 8pt 14pt; border-radius: 6pt;
      flex: 1;
    }
    .stat .num {
      display: block; font-family: 'DM Sans'; font-size: 18pt;
      font-weight: 600; color: var(--ocean-deep);
    }
    .stat .label {
      font-size: 9pt; color: var(--muted); text-transform: uppercase;
      letter-spacing: 0.06em;
    }
    table {
      width: 100%; border-collapse: collapse;
      font-size: 10pt; margin: 8pt 0;
    }
    th {
      background: linear-gradient(135deg, var(--ocean-mid), var(--ocean-shallow));
      color: white; text-align: left; padding: 6pt 8pt;
      font-family: 'DM Sans', system-ui, sans-serif;
    }
    td { padding: 4pt 8pt; border-bottom: 1px solid var(--mist); }
    tr:nth-child(even) td { background: rgba(232, 241, 248, 0.4); }
    .pill {
      display: inline-block; padding: 1pt 8pt; border-radius: 12pt;
      font-size: 9pt; font-weight: 600; color: white;
    }
    .pill-reinforcing { background: var(--ocean-surface); }
    .pill-balancing { background: #dc131e; }
    footer {
      margin-top: 32pt; padding-top: 8pt;
      border-top: 1px solid var(--mist);
      font-size: 9pt; color: var(--muted); text-align: center;
    }
  </style>
</head>
<body>

<h1>{{ project.metadata.name }}</h1>
<p class="meta">
  SES Toolbox Report &nbsp;·&nbsp;
  Generated {{ generated_at }}
  {% if project.metadata.modified_at %}
    &nbsp;·&nbsp; Last modified {{ project.metadata.modified_at }}
  {% endif %}
</p>

<h2>Summary</h2>
<div class="summary">
  <div class="stat"><span class="num">{{ summary.elements }}</span>
    <span class="label">Elements</span></div>
  <div class="stat"><span class="num">{{ summary.connections }}</span>
    <span class="label">Connections</span></div>
  <div class="stat"><span class="num">{{ summary.loops }}</span>
    <span class="label">Feedback loops</span></div>
  <div class="stat"><span class="num">{{ summary.density }}</span>
    <span class="label">Density</span></div>
</div>

<h3>Element types</h3>
<table>
  <thead><tr><th>DAPSIWRM type</th><th style="text-align:right">Count</th></tr></thead>
  <tbody>
  {% for type, count in summary.types_by_count %}
    <tr><td>{{ type }}</td><td style="text-align:right">{{ count }}</td></tr>
  {% endfor %}
  </tbody>
</table>

{% if loops %}
<h2>Feedback loops</h2>
<p class="meta">{{ summary.reinforcing }} reinforcing, {{ summary.balancing }} balancing.</p>
<table>
  <thead><tr><th>ID</th><th>Type</th><th>Length</th><th>Path</th></tr></thead>
  <tbody>
  {% for cycle in loops[:20] %}
    <tr>
      <td>{{ cycle.id }}</td>
      <td><span class="pill pill-{{ cycle.type|lower }}">{{ cycle.type }}</span></td>
      <td>{{ cycle.length }}</td>
      <td style="font-size: 9pt;">{{ cycle.path }}</td>
    </tr>
  {% endfor %}
  </tbody>
</table>
{% if loops|length > 20 %}
<p class="meta">… and {{ loops|length - 20 }} more (truncated for the report).</p>
{% endif %}
{% endif %}

<h2>Top centrality</h2>
{% for metric, rows in top_centrality.items() %}
<h3>{{ metric|capitalize }}</h3>
<table>
  <thead><tr><th>Rank</th><th>Element</th><th>Type</th>
    <th style="text-align:right">Score</th></tr></thead>
  <tbody>
  {% for r in rows %}
    <tr><td>{{ r.rank }}</td><td>{{ r.label }}</td><td>{{ r.type }}</td>
        <td style="text-align:right">{{ '%.4f'|format(r.value) }}</td></tr>
  {% endfor %}
  </tbody>
</table>
{% endfor %}

<h2>Leverage points</h2>
<p class="meta">Composite score = standardised betweenness + eigenvector + PageRank.</p>
<table>
  <thead><tr><th>Rank</th><th>Element</th><th>Type</th>
    <th style="text-align:right">Score</th></tr></thead>
  <tbody>
  {% for r in leverage[:10] %}
    <tr><td>{{ r.rank }}</td><td>{{ r.label }}</td><td>{{ r.type }}</td>
        <td style="text-align:right">{{ '%.3f'|format(r.score) }}</td></tr>
  {% endfor %}
  </tbody>
</table>

<footer>
  MarineSABRES SES Toolbox &nbsp;·&nbsp; Horizon Europe Project
</footer>

</body>
</html>
"""


def _build_context(project: Project) -> dict[str, Any]:
    isa = project.isa_data

    # Summary stats
    elements_by_type: dict[str, int] = {}
    for el in isa.elements:
        elements_by_type[el.type] = elements_by_type.get(el.type, 0) + 1

    metrics = net_analysis.basic_metrics(isa)
    cycles = net_analysis.feedback_loops(isa)
    loops = net_analysis.classify_loops(cycles, isa)
    reinforcing = sum(1 for r in loops if r["type"] == "Reinforcing")
    balancing = len(loops) - reinforcing

    # Top centrality per metric (top-5)
    top_centrality = {
        m: net_analysis.top_n_by_metric(isa, m, n=5)
        for m in net_analysis.CENTRALITY_METRICS
    }

    # Leverage ranking
    lev_scores = net_analysis.leverage_scores(isa)
    by_id = {el.id: el for el in isa.elements}
    leverage_rows = sorted(lev_scores.items(), key=lambda kv: kv[1], reverse=True)
    leverage = [
        {
            "rank": rank,
            "id": nid,
            "label": (by_id[nid].label if nid in by_id else nid),
            "type":  (by_id[nid].type  if nid in by_id else ""),
            "score": value,
        }
        for rank, (nid, value) in enumerate(leverage_rows, start=1)
    ]

    return {
        "project": project,
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "summary": {
            "elements": metrics["nodes"],
            "connections": metrics["edges"],
            "loops": len(loops),
            "density": f"{metrics['density']:.4f}",
            "types_by_count": sorted(elements_by_type.items(),
                                     key=lambda kv: kv[1], reverse=True),
            "reinforcing": reinforcing,
            "balancing": balancing,
        },
        "loops": loops,
        "top_centrality": top_centrality,
        "leverage": leverage,
    }


def render_html(project: Project) -> str:
    """Render the report as a self-contained HTML string."""
    env = Environment(loader=BaseLoader(), autoescape=select_autoescape(["html"]))
    template = env.from_string(_TEMPLATE)
    return template.render(**_build_context(project))


def render_pdf(project: Project) -> bytes:
    """Render the report as PDF bytes (WeasyPrint underneath)."""
    from weasyprint import HTML

    html = render_html(project)
    return HTML(string=html).write_pdf()


def render_docx(project: Project) -> bytes:
    """Render the report as a Word .docx, returned as bytes ready for
    `@render.download` to stream. Uses `python-docx` — same content as
    the HTML/PDF, restructured into native Word elements (heading styles,
    tables) so the file edits cleanly in Word.

    R counterpart: the `officer` + `flextable` pipeline used in
    `functions/export_functions.R`. python-docx covers ~90% of officer's
    surface; the remaining 10% (advanced styling) we don't currently need.
    """
    from io import BytesIO

    from docx import Document
    from docx.shared import Pt, RGBColor

    ctx = _build_context(project)
    summary = ctx["summary"]
    doc = Document()

    # ---- Title + meta -----------------------------------------------------
    title = doc.add_heading(project.metadata.name, level=0)
    title.style.font.size = Pt(22)
    meta = doc.add_paragraph()
    meta.add_run(
        f"SES Toolbox Report  ·  Generated {ctx['generated_at']}"
    ).italic = True
    if project.metadata.modified_at:
        meta.add_run(f"  ·  Last modified {project.metadata.modified_at}").italic = True

    # ---- Summary ----------------------------------------------------------
    doc.add_heading("Summary", level=1)
    bullet_lines = [
        f"Elements:    {summary['elements']}",
        f"Connections: {summary['connections']}",
        f"Feedback loops: {summary['loops']}  "
        f"({summary['reinforcing']} reinforcing, {summary['balancing']} balancing)",
        f"Density:     {summary['density']}",
    ]
    for line in bullet_lines:
        doc.add_paragraph(line, style="List Bullet")

    # ---- Element types breakdown -----------------------------------------
    doc.add_heading("Element types", level=2)
    type_table = doc.add_table(rows=1, cols=2)
    type_table.style = "Light Grid Accent 1"
    hdr = type_table.rows[0].cells
    hdr[0].text, hdr[1].text = "DAPSIWRM type", "Count"
    for type_name, count in summary["types_by_count"]:
        row = type_table.add_row().cells
        row[0].text = type_name
        row[1].text = str(count)

    # ---- Feedback loops ---------------------------------------------------
    if ctx["loops"]:
        doc.add_heading("Feedback loops", level=1)
        loops_table = doc.add_table(rows=1, cols=4)
        loops_table.style = "Light Grid Accent 1"
        hdr = loops_table.rows[0].cells
        hdr[0].text, hdr[1].text = "ID", "Type"
        hdr[2].text, hdr[3].text = "Length", "Path"
        for cycle in ctx["loops"][:20]:
            row = loops_table.add_row().cells
            row[0].text = cycle["id"]
            row[1].text = cycle["type"]
            row[2].text = str(cycle["length"])
            row[3].text = cycle["path"]
        if len(ctx["loops"]) > 20:
            doc.add_paragraph(
                f"… and {len(ctx['loops']) - 20} more (truncated)."
            ).italic = True

    # ---- Top centrality ---------------------------------------------------
    doc.add_heading("Top centrality", level=1)
    for metric, rows in ctx["top_centrality"].items():
        doc.add_heading(metric.title(), level=2)
        ct = doc.add_table(rows=1, cols=4)
        ct.style = "Light Grid Accent 1"
        hdr = ct.rows[0].cells
        hdr[0].text, hdr[1].text = "Rank", "Element"
        hdr[2].text, hdr[3].text = "Type", "Score"
        for r in rows:
            row = ct.add_row().cells
            row[0].text = str(r["rank"])
            row[1].text = r["label"]
            row[2].text = r["type"]
            row[3].text = f"{r['value']:.4f}"

    # ---- Leverage ---------------------------------------------------------
    doc.add_heading("Leverage points", level=1)
    doc.add_paragraph(
        "Composite score = standardised betweenness + eigenvector + PageRank."
    ).italic = True
    lev_table = doc.add_table(rows=1, cols=4)
    lev_table.style = "Light Grid Accent 1"
    hdr = lev_table.rows[0].cells
    hdr[0].text, hdr[1].text = "Rank", "Element"
    hdr[2].text, hdr[3].text = "Type", "Score"
    for r in ctx["leverage"][:10]:
        row = lev_table.add_row().cells
        row[0].text = str(r["rank"])
        row[1].text = r["label"]
        row[2].text = r["type"]
        row[3].text = f"{r['score']:.3f}"

    # ---- Footer-ish line --------------------------------------------------
    doc.add_paragraph()
    footer = doc.add_paragraph(
        "MarineSABRES SES Toolbox  ·  Horizon Europe Project"
    )
    footer.runs[0].italic = True
    footer.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
